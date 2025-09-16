# PDF parsing, chunking
import PyPDF2
import re
from typing import List, Dict, Any
from datetime import datetime
from pathlib import Path
from docx import Document as DocxDocument


def process_file(file_path):
    """Process file based on its type and extract content with metadata"""
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Detect file type and extract content
    if path.suffix.lower() == ".pdf":
        return _extract_pdf(file_path)
    elif path.suffix.lower() == ".docx":
        return _extract_docx(file_path)
    elif path.suffix.lower() in ['.txt', '.md']:
        return _extract_text(file_path)
    else:
        raise ValueError(f"Unsupported file format: {path.suffix}")


def _extract_pdf(file_path):
    """Extracting details from the pdf file format including the metadata"""
    pages_data = extract_text_with_pages(file_path)
    meta_data = extract_pdf_metadata(file_path)
    
    # Chunk with page metadata preserved
    chunks = chunk_pages_with_metadata(pages_data, meta_data)
    
    return {'content': chunks, 'metadata': meta_data}


def extract_text_with_pages(file_path: str):
    """Extract text preserving page boundaries"""
    path = Path(file_path)
    content = []
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                content.append({'page_number': page_num+1,
                                'text': page_text})
        return content
    except Exception as e:
        raise Exception(f"Error processing PDF {file_path}: {str(e)}")


def extract_pdf_metadata(file_path):
    """Extract document-level metadata"""
    path = Path(file_path)
    
    try:
        file_stats = path.stat()
        pdf_reader = PyPDF2.PdfReader(file_path)
        
        # Safely extract PDF metadata with fallbacks
        pdf_metadata = pdf_reader.metadata or {}
        
        return {
            'filename': path.name,
            'file_size': file_stats.st_size,
            'file_type': path.suffix.lower(),
            'page_count': len(pdf_reader.pages),
            'file_title': pdf_metadata.get('title', '') or '',
            'author': pdf_metadata.get('author', '') or '',
            'subject': pdf_metadata.get('subject', '') or '',
            'creator': pdf_metadata.get('creator', '') or '',
            'creation_date': pdf_metadata.get('creation_date', '') or ''
        }
    except Exception as e:
        raise Exception(f"Error extracting metadata from {file_path}: {str(e)}")


def smart_chunking(text, target_size=150, overlap_size=30):
    """Chunking with pre-defined chunk size and overlap between the two chunks"""
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if current_chunk:
            potential_chunk = current_chunk + " " + sentence
        else:
            potential_chunk = sentence
        
        if len(potential_chunk) <= target_size:  
            current_chunk = potential_chunk
        else: 
            if current_chunk: 
                chunks.append({
                    'content': current_chunk.strip(),
                    'length': len(current_chunk),
                })
            
            # Creating overlap for the next chunk
            if len(current_chunk) >= overlap_size:
                overlap = current_chunk[-overlap_size:]
            else:
                overlap = current_chunk
            
            # Start new chunk with overlap and current sentence
            if overlap:
                current_chunk = overlap + " " + sentence
            else:
                current_chunk = sentence
            
    # Adding final chunk
    if current_chunk.strip():
        chunks.append({
            'content': current_chunk.strip(),
            'length': len(current_chunk),
        })
    
    return chunks


def chunk_pages_with_metadata(pages_data, doc_metadata, target_size=150, overlap_size=30):
    """Chunk pages while preserving page numbers"""
    all_chunks = []
    
    for page in pages_data:
        # Using smart chunking
        page_chunks = smart_chunking(page['text'], target_size, overlap_size)
        
        # Add page number and document info to each chunk
        for chunk in page_chunks:
            chunk['page_number'] = page['page_number']
            chunk['doc_id'] = doc_metadata['filename']
            all_chunks.append(chunk)
    
    return all_chunks


def _extract_docx(file_path):
    """Extract content from DOCX file"""
    path = Path(file_path)

    try:
        doc = DocxDocument(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        file_stats = path.stat()

        return {
            'content': content,
            'metadata': {
                'filename': path.name,
                'file_size': file_stats.st_size,
                'file_type': path.suffix.lower(),
                'paragraph_count': len(doc.paragraphs),
                'character_count': len(content),
                'title': doc.core_properties.title or ''
            }
        }
    except Exception as e:
        raise Exception(f"Error processing DOCX {file_path}: {str(e)}")


def _extract_text(file_path: str) -> Dict[str, Any]:
    """Extract content from text files"""
    path = Path(file_path)
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        file_stats = path.stat()
        
        return {
            'content': content,
            'metadata': {
                'filename': path.name,
                'file_size': file_stats.st_size,
                'file_type': path.suffix.lower(),
                'character_count': len(content)
            }
        }
    except Exception as e:
        raise Exception(f"Error processing text file {file_path}: {str(e)}")